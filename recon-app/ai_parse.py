"""
Document text extraction + Anthropic Claude AI analysis for PSP agreements.

The FEES system schema this module targets:
  psp_agreements  — one row per PSP
  psp_fee_rules   — many rows per PSP, each a specific fee line
  psp_fee_tiers   — volume bands for tiered-rate rules
"""
import io
import os
import json
import anthropic

ANTHROPIC_KEY   = os.environ.get("ANTHROPIC_API_KEY", "")
ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6")

# Head+tail strategy: 4 k of agreement header + 12 k of tail (where fee schedules live)
HEAD_CHARS = 4_000
TAIL_CHARS = 12_000

# ── Canonical value lists — MUST match app.py exactly ─────────────────────
FEE_TYPES = [
    "Deposit", "Withdrawal", "Settlement", "Chargeback", "Refund",
    "Rolling Reserve", "Holdback", "Setup", "Registration", "Minimum Monthly",
]
PAYMENT_METHODS = [
    "Credit Cards", "Bank Wire", "Mobile Money", "Electronic Payment",
    "Crypto", "MOMO", "E-Wallet",
]

# ── System prompt ──────────────────────────────────────────────────────────
SYSTEM_PROMPT = """\
You are a specialist financial-document parser for a PSP (Payment Service Provider) \
reconciliation system called FEES. Your sole job is to extract structured data from \
PSP merchant agreements and return it as a single, valid JSON object that exactly \
matches the schema shown below. No markdown fences, no explanation, no extra keys.

═══════════════════════════════════════════════════════════
  OUTPUT JSON SCHEMA  (use null for any unknown field)
═══════════════════════════════════════════════════════════
{
  "agreement": {
    "psp_name"         : "<SHORT UPPERCASE NAME — e.g. NUVEI, BITOLO>",
    "provider_name"    : "<Full legal company name of the PSP — e.g. NUVEI LIMITED>",
    "agreement_entity" : "<CMT legal entity that signed — see rules below>",
    "agreement_date"   : "<YYYY-MM-DD or null>",
    "addendum_date"    : "<YYYY-MM-DD or null — most recent addendum/amendment date>",
    "auto_settlement"  : <true | false>,
    "settlement_bank"  : "<Bank name + account identifier, or null>"
  },
  "fee_rules": [
    {
      "payment_method" : "<one value from the PAYMENT METHODS list, or null>",
      "fee_type"       : "<one value from the FEE TYPES list>",
      "country"        : "<country name as written in agreement, or 'GLOBAL'>",
      "sub_provider"   : "<card network / mobile operator / sub-brand, or null>",
      "fee_kind"       : "<'percentage' | 'fixed' | 'fixed_plus_pct' | 'tiered'>",
      "pct_rate"       : <decimal fraction or null — e.g. 3.5 % → 0.035>,
      "fixed_amount"   : <number or null>,
      "fixed_currency" : "<ISO-4217 code or null — e.g. USD, EUR, NGN>",
      "description"    : "<verbatim or paraphrased fee label from the document, or null>",
      "tiers"          : [
        { "volume_from": <number>, "volume_to": <number or null>, "pct_rate": <decimal fraction> }
      ]
    }
  ]
}

═══════════════════════════════════════════════════════════
  FIELD-BY-FIELD RULES
═══════════════════════════════════════════════════════════

── agreement.psp_name ──────────────────────────────────
• Short commercial name (all caps), NOT the legal entity name.
  ✓ "NUVEI"   ✓ "BITOLO"   ✓ "DIRECTA24"   ✗ "NUVEI LIMITED"

── agreement.provider_name ─────────────────────────────
• Full legal name of the PSP company as it appears in the agreement.
  ✓ "NUVEI LIMITED"   ✓ "FLYBRIDGE LIMITED"   ✓ "BLAVEN TECHNOLOGIES LLP"

── agreement.agreement_entity ──────────────────────────
• The CMT Group legal entity that signed (counterparty).
• Common values (normalise spelling to these exactly):
    "CMT PROCESSING LTD"    ← formerly "CMT PROCESSING LIMITED"
    "GCMT GROUP LTD"        ← also appears as "GCMT GROUP LIMITED"
• If neither is clear, use null.

── agreement.agreement_date / addendum_date ────────────
• Convert any date format to YYYY-MM-DD:
    "27.10.2025" → "2025-10-27"
    "28/06/2022" → "2022-06-28"
    "January 15, 2024" → "2024-01-15"
    "n/a", "N/A", "n.a." → null

── agreement.auto_settlement ───────────────────────────
• true  → document mentions "automatic settlement", "auto settlement", "YES" in settlement column,
          or says the PSP remits funds automatically on a fixed schedule.
• false → manual settlement, "NO", or no mention.

── agreement.settlement_bank ───────────────────────────
• Capture the full string including bank name and account/IBAN reference.
  e.g. "Alpha CY77009005150005151060573155"
  or   "Absa 242586-USD-1046-11"

── fee_rules[].payment_method ──────────────────────────
ALLOWED VALUES (map document terms to these exactly):
  "Credit Cards"        ← CREDIT CARDS, CARD, VISA, MASTERCARD, AMEX
  "Bank Wire"           ← BANK WIRE, WIRE TRANSFER, WIRE, BANK TRANSFER
  "Mobile Money"        ← MOBILE MONEY, MOMO, MOBILE PAYMENTS
  "Electronic Payment"  ← ELECTRONIC PAYMENT, E-WALLET (general), NETELLER, SKRILL,
                          PAYSAFE, PAYPAL, ONLINE BANKING
  "Crypto"              ← CRYPTO, CRYPTOCURRENCY, CRYPTOWALLET, BITCOIN, USDT
  "MOMO"                ← use only when document explicitly labels it MOMO
  "E-Wallet"            ← use only when document labels it E-WALLET distinctly from Electronic Payment
  null                  ← when a fee applies to all methods, or method is unspecified

── fee_rules[].fee_type ────────────────────────────────
ALLOWED VALUES (map document terms to these exactly):
  "Deposit"          ← DEPOSIT FEE, ACQUIRING FEE, PROCESSING FEE (incoming), DEPOSIT AND WITHDRAWAL FEE*
  "Withdrawal"       ← WITHDRAWAL FEE, CASHOUT FEE, PAYOUT FEE, DEPOSIT AND WITHDRAWAL FEE*
  "Settlement"       ← SETTLEMENT FEE, SETTLEMENT CHARGE
  "Chargeback"       ← CHARGEBACK FEE, DISPUTE FEE
  "Refund"           ← REFUND FEE, REVERSAL FEE
  "Rolling Reserve"  ← ROLLING RESERVE, RESERVE RATE, HOLDBACK RATE (as reserve, not one-off)
  "Holdback"         ← HOLDBACK (one-off amount held back on termination), HOLDBACK AMOUNT
  "Setup"            ← SETUP FEE, SET UP FEE, ONBOARDING FEE, INTEGRATION FEE
  "Registration"     ← REGISTRATION FEE, ANNUAL CARD SCHEME FEE, CARD BRAND FEE
  "Minimum Monthly"  ← MINIMUM FEE, MINIMUM MONTHLY FEE, MONTHLY MINIMUM
  * "DEPOSIT AND WITHDRAWAL FEE" — create TWO separate rules, one Deposit and one Withdrawal,
    each with the same rate, unless rates differ.

── fee_rules[].country ─────────────────────────────────
• Use the country name exactly as written in the document.
• If the fee applies globally / to all countries → "GLOBAL"
• Multiple countries in one row → create one rule per country.
• "ALL" or "ALL COUNTRIES" or no country specified → "GLOBAL"

── fee_rules[].sub_provider ────────────────────────────
• Card networks: "Visa", "Mastercard", "Amex", "UnionPay"
• Mobile operators: "MTN", "Airtel", "Vodacom", "M-Pesa", "Orange"
• If the fee applies to all sub-providers → null

── fee_rules[].fee_kind and rate fields ─────────────────
This is the most critical section. Apply the following priority order:

  STEP 1 — Look for explicit % sign:
    "3.5%"   → fee_kind="percentage", pct_rate=0.035
    "3,5%"   → European comma decimal; same result: pct_rate=0.035
    "0,5%"   → pct_rate=0.005
    "1,1%"   → pct_rate=0.011

  STEP 2 — Decimal fraction < 1.0 with no % and no currency code:
    These are stored as decimal fractions in the source system:
    "0.013"  → fee_kind="percentage", pct_rate=0.013   (= 1.3%)
    "0.02"   → fee_kind="percentage", pct_rate=0.02    (= 2%)
    "0.03"   → fee_kind="percentage", pct_rate=0.03    (= 3%)
    "0.29"   — AMBIGUOUS. Check fee_type:
               • If fee_type is Deposit/Withdrawal/Settlement → likely percentage, pct_rate=0.29
               • If combined with "per transaction" language  → likely fixed, fixed_amount=0.29, fixed_currency="USD"

  STEP 3 — Integer or large decimal with no % and no currency code:
    These are usually fixed amounts in USD (unless context implies percentage):
    "0.8"    with "per transaction" → fixed, fixed_amount=0.8, fixed_currency="USD"
    "1"      as settlement fee → fee_kind="percentage", pct_rate=0.01  (= 1%, common for settlement)
    "1"      as a per-transaction fee → fixed, fixed_amount=1, fixed_currency="USD"
    "5"      as margin → fixed, fixed_amount=5, fixed_currency="USD" (or EUR if specified)
    "18"     as chargeback exceeded fee → fixed, fixed_amount=18, fixed_currency="USD"
    "100"    as dormancy/inactivity → fixed, fixed_amount=100, fixed_currency="USD"
    "500"    for Setup/Registration → fixed, fixed_amount=500, fixed_currency="USD"
    Use the fee_type and surrounding description text as context to disambiguate.

  STEP 4 — Combined fixed + percentage:
    "100 NGN + 1%"   → fee_kind="fixed_plus_pct", fixed_amount=100, fixed_currency="NGN", pct_rate=0.01
    "8 GHS + 1%"     → fee_kind="fixed_plus_pct", fixed_amount=8,   fixed_currency="GHS", pct_rate=0.01
    "€0.30 + 2.9%"   → fee_kind="fixed_plus_pct", fixed_amount=0.30, fixed_currency="EUR", pct_rate=0.029
    "USD 0.50 + 3%"  → fee_kind="fixed_plus_pct", fixed_amount=0.50, fixed_currency="USD", pct_rate=0.03

  STEP 5 — Tiered / volume-based:
    Document shows a table with volume bands and different rates per band.
    fee_kind="tiered", pct_rate=null, fixed_amount=null
    Populate tiers array with { volume_from, volume_to, pct_rate } for each band.
    volume_to = null for the last (open-ended) tier.

  NEVER invent a rate. If genuinely ambiguous, choose the most contextually likely
  interpretation and note it in the description field.

── fee_rules[].description ─────────────────────────────
• Copy the exact fee label or description from the document.
• Keep it concise (≤ 120 chars). Include any qualifiers (e.g. "on approved transactions",
  "per merchant ID", "max 30 USD", "where applicable").

═══════════════════════════════════════════════════════════
  WORKED EXAMPLE (from a real agreement)
═══════════════════════════════════════════════════════════
Source row: NUVEI | CREDIT CARDS | DEPOSIT FEE | GLOBAL | ACQUIRING FEE | 0.01 | YES | Alpha CY77...

Output rule:
{
  "payment_method": "Credit Cards",
  "fee_type": "Deposit",
  "country": "GLOBAL",
  "sub_provider": null,
  "fee_kind": "percentage",
  "pct_rate": 0.01,
  "fixed_amount": null,
  "fixed_currency": null,
  "description": "Acquiring fee - on approved transactions",
  "tiers": []
}

Source row: BITOLO | BANK WIRE | WITHDRAWAL FEE | NIGERIA | 100 NGN + 1%

Output rule:
{
  "payment_method": "Bank Wire",
  "fee_type": "Withdrawal",
  "country": "Nigeria",
  "sub_provider": null,
  "fee_kind": "fixed_plus_pct",
  "pct_rate": 0.01,
  "fixed_amount": 100,
  "fixed_currency": "NGN",
  "description": "Withdrawal fee",
  "tiers": []
}

Source row: NUVEI | CREDIT CARDS | SET UP FEE | GLOBAL | PER MERCHANT ID | 500

Output rule:
{
  "payment_method": "Credit Cards",
  "fee_type": "Setup",
  "country": "GLOBAL",
  "sub_provider": null,
  "fee_kind": "fixed",
  "pct_rate": null,
  "fixed_amount": 500,
  "fixed_currency": "USD",
  "description": "Per merchant ID",
  "tiers": []
}

═══════════════════════════════════════════════════════════
  EXTRACTION STRATEGY
═══════════════════════════════════════════════════════════
• Scan the ENTIRE document. Fee schedules are frequently in appendices,
  schedules, or exhibit sections near the END of the document.
• Look for: tables, bullet lists, sections titled "Fees", "Pricing",
  "Schedule", "Charges", "Commercial Terms", "Fee Schedule", "Exhibit".
• Extract EVERY distinct fee line as a separate entry in fee_rules.
• If a fee applies to multiple countries listed separately, create one
  rule per country.
• Skip non-fee items: KYC requirements, legal boilerplate, liability
  clauses, privacy policy, data protection terms.
• If no fee schedule is found, return fee_rules as an empty array [].
"""

# ── JSON schema shown to model in the user turn ───────────────────────────
_JSON_EXAMPLE = json.dumps({
    "agreement": {
        "psp_name": "EXAMPLE",
        "provider_name": "Example Ltd",
        "agreement_entity": "CMT PROCESSING LTD",
        "agreement_date": "2024-01-15",
        "addendum_date": None,
        "auto_settlement": True,
        "settlement_bank": "Alpha CY77009005150005151060573155"
    },
    "fee_rules": [
        {
            "payment_method": "Credit Cards",
            "fee_type": "Deposit",
            "country": "GLOBAL",
            "sub_provider": None,
            "fee_kind": "percentage",
            "pct_rate": 0.035,
            "fixed_amount": None,
            "fixed_currency": None,
            "description": "Processing fee on approved transactions",
            "tiers": []
        }
    ]
}, indent=2)


# ── JSON Schema (used when model supports structured outputs) ──────────────
# This is the strict contract the API enforces at the token level when using
# response_format={"type":"json_schema"}.  When that mode is unavailable the
# schema is still shown to the model in the user prompt as a reference.
_OUTPUT_SCHEMA = {
    "name":   "psp_agreement_extraction",
    "strict": True,
    "schema": {
        "type": "object",
        "required": ["agreement", "fee_rules"],
        "additionalProperties": False,
        "properties": {
            "agreement": {
                "type": "object",
                "required": [
                    "psp_name", "provider_name", "agreement_entity",
                    "agreement_date", "addendum_date",
                    "auto_settlement", "settlement_bank",
                ],
                "additionalProperties": False,
                "properties": {
                    "psp_name":          {"type": "string"},
                    "provider_name":     {"type": ["string", "null"]},
                    "agreement_entity":  {"type": ["string", "null"]},
                    "agreement_date":    {"type": ["string", "null"]},
                    "addendum_date":     {"type": ["string", "null"]},
                    "auto_settlement":   {"type": "boolean"},
                    "settlement_bank":   {"type": ["string", "null"]},
                },
            },
            "fee_rules": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": [
                        "payment_method", "fee_type", "country", "sub_provider",
                        "fee_kind", "pct_rate", "fixed_amount", "fixed_currency",
                        "description", "tiers",
                    ],
                    "additionalProperties": False,
                    "properties": {
                        "payment_method":  {"type": ["string", "null"], "enum": PAYMENT_METHODS + [None]},
                        "fee_type":        {"type": "string", "enum": FEE_TYPES},
                        "country":         {"type": "string"},
                        "sub_provider":    {"type": ["string", "null"]},
                        "fee_kind":        {"type": "string", "enum": ["percentage", "fixed", "fixed_plus_pct", "tiered"]},
                        "pct_rate":        {"type": ["number", "null"]},
                        "fixed_amount":    {"type": ["number", "null"]},
                        "fixed_currency":  {"type": ["string", "null"]},
                        "description":     {"type": ["string", "null"]},
                        "tiers": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "required": ["volume_from", "volume_to", "pct_rate"],
                                "additionalProperties": False,
                                "properties": {
                                    "volume_from": {"type": "number"},
                                    "volume_to":   {"type": ["number", "null"]},
                                    "pct_rate":    {"type": "number"},
                                },
                            },
                        },
                    },
                },
            },
        },
    },
}


# ── Text extraction ────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a PDF or DOCX file."""
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)

    if ext in ("docx", "doc"):
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Tables are critical — fee schedules are often formatted as tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise ValueError(f"Unsupported file type: .{ext}  (only PDF and DOCX are supported)")


def extract_pages(file_bytes: bytes, filename: str) -> list:
    """Return [(page_num, text), ...] for source-location lookup.
    Page numbers are 1-indexed. DOCX returns [(1, full_text)].

    For PDFs, tables are re-extracted with extract_tables() and appended
    as pipe-separated rows so snippets remain human-readable.
    pdfplumber extract_text() flattens table cells into a string like
    "4.4% 4.4% 4.4% 2.25%" which is unreadable as a source quote.
    """
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                parts = []
                plain = page.extract_text()
                if plain and plain.strip():
                    parts.append(plain)
                for table in (page.extract_tables() or []):
                    for row in table:
                        cells = [str(c or "").strip() for c in row]
                        line  = " | ".join(cells)
                        if line.strip(" |"):
                            parts.append(line)
                combined = chr(10).join(parts)
                if combined.strip():
                    pages.append((i, combined))
        return pages
    if ext in ("docx", "doc"):
        return [(1, extract_text(file_bytes, filename))]
    return []

def annotate_sources(rules: list, pages: list) -> None:
    """
    Mutate each rule dict in-place, adding:
      source_page        : int | None    -- 1-indexed PDF page
      source_quote       : str | None    -- verbatim snippet <= 160 chars
      source_confidence  : float | None  -- 0.0-1.0 match probability

    Probabilistic scoring model
    ---------------------------
    For every occurrence of the rate number on a page compute three
    proximity sub-scores (each 0-1) using inverse-distance decay bands:

      pct_score     -- distance to nearest '%'
                       (0-4 -> 1.0, 5-20 -> 0.8, 21-60 -> 0.55,
                        61-150 -> 0.30, 151-400 -> 0.10, beyond -> 0.0)
      country_score -- distance to country name
                       (0-60 -> 1.0, 61-200 -> 0.70, 201-450 -> 0.35, beyond -> 0.05)
      method_score  -- distance to payment method word
                       (0-120 -> 1.0, 121-400 -> 0.60, 401-800 -> 0.25, beyond -> 0.10)

    Combined probability = pct*0.55 + country*0.30 + method*0.15

    % proximity is weighted most heavily: a number far from any % is
    likely a registration number, page reference, or section heading.
    Results with probability < 0.30 are suppressed (shown as None).
    """
    import re

    MIN_CONFIDENCE = 0.30

    def _decay(dist, bands):
        for max_d, s in bands:
            if dist <= max_d:
                return s
        return bands[-1][1]

    PCT_BANDS     = [(4, 1.0), (20, 0.80), (60, 0.55), (150, 0.30), (400, 0.10), (9999, 0.0)]
    COUNTRY_BANDS = [(60, 1.0), (200, 0.70), (450, 0.35), (9999, 0.05)]
    METHOD_BANDS  = [(120, 1.0), (400, 0.60), (800, 0.25), (9999, 0.10)]

    def _rate_vals(rule):
        vals = []
        pct = rule.get("pct_rate")
        if pct is not None:
            v = pct * 100
            for fmt in (f"{v:.4g}", f"{v:.2f}", f"{v:.1f}"):
                if fmt not in vals:
                    vals.append(fmt)
        for tier in rule.get("tiers") or []:
            tr = tier.get("pct_rate")
            if tr is not None:
                v2 = f"{tr*100:.4g}"
                if v2 not in vals:
                    vals.append(v2)
        return vals

    for rule in rules:
        country = (rule.get("country") or "").lower().strip()
        if country == "global":
            country = ""
        method  = (rule.get("payment_method") or "").lower()
        vals    = _rate_vals(rule)

        best_page  = None
        best_quote = None
        best_conf  = 0.0

        for page_num, page_text in pages:
            tl = page_text.lower()

            pct_pos     = [m.start() for m in re.finditer(r'%', tl)]
            country_pos = ([m.start() for m in re.finditer(re.escape(country), tl)]
                           if country else [])
            method_words = [w for w in re.split(r'\W+', method) if len(w) > 3]
            method_pos  = []
            for w in method_words:
                method_pos += [m.start() for m in re.finditer(re.escape(w), tl)]

            page_best_conf = 0.0
            page_best_idx  = -1

            for val in vals:
                pat = re.compile(r'(?<!\d)' + re.escape(val) + r'(?!\d)')
                for m in pat.finditer(tl):
                    pos = m.start()

                    d_pct = min((abs(pos - p) for p in pct_pos), default=9999)
                    ps = _decay(d_pct, PCT_BANDS)

                    if country_pos:
                        d_co = min(abs(pos - p) for p in country_pos)
                        cs = _decay(d_co, COUNTRY_BANDS)
                    else:
                        cs = 0.5

                    if method_pos:
                        d_me = min(abs(pos - p) for p in method_pos)
                        ms = _decay(d_me, METHOD_BANDS)
                    else:
                        ms = 0.2

                    conf = ps * 0.55 + cs * 0.30 + ms * 0.15

                    if conf > page_best_conf:
                        page_best_conf = conf
                        page_best_idx  = pos

            if page_best_conf > best_conf:
                best_conf  = page_best_conf
                best_page  = page_num
                best_quote = None
                if page_best_idx >= 0:
                    lo = page_best_idx
                    hi = page_best_idx + 6
                    if country_pos:
                        nearest_co = min(country_pos, key=lambda p: abs(p - page_best_idx))
                        if abs(nearest_co - page_best_idx) <= 450:
                            lo = min(lo, nearest_co)
                            hi = max(hi, nearest_co + len(country))
                    start = max(0, lo - 25)
                    end   = min(len(page_text), hi + 130)
                    best_quote = re.sub(r'\s+', ' ', page_text[start:end]).strip()

        rule["source_page"]       = best_page  if best_conf >= MIN_CONFIDENCE else None
        rule["source_quote"]      = best_quote if best_conf >= MIN_CONFIDENCE else None
        rule["source_confidence"] = round(best_conf, 2) if best_conf >= MIN_CONFIDENCE else None



def find_potential_gaps(rules: list, pages: list) -> list:
    """
    Heuristic scan for fee-like percentage values in the document that
    were not matched to any extracted rule.

    Strategy:
    - For every page that has at least one matched rule, count all
      "fee-like" % occurrences (0.1 – 30%, word-boundary matched).
    - If the count significantly exceeds the number of matched rules on
      that page, flag the page as potentially having missed entries and
      return a sample snippet for review.
    - Also flags pages with any fee-like % that have NO matched rules
      at all, if at least one other page does have matches (i.e. the AI
      found some fee schedule pages but skipped others).

    Returns list of dicts: {page, pct_found, rules_matched, snippet}.
    Capped at 8 items to avoid noise.
    """
    import re

    fee_pct = re.compile(r'(?<!\d)(\d{1,2}\.?\d{0,2})\s{0,4}%')

    # Count matched rules per page (only pages with confident matches)
    rules_by_page = {}
    for rule in rules:
        p = rule.get("source_page")
        if p and (rule.get("source_confidence") or 0) >= 0.30:
            rules_by_page[p] = rules_by_page.get(p, 0) + 1

    any_match_pages = set(rules_by_page)

    gaps = []
    for page_num, page_text in pages:
        tl = page_text.lower()
        fee_hits = []
        for m in fee_pct.finditer(tl):
            try:
                v = float(m.group(1))
            except ValueError:
                continue
            if 0.1 <= v <= 30:          # plausible fee range
                fee_hits.append((m.start(), m.group()))

        if not fee_hits:
            continue

        matched = rules_by_page.get(page_num, 0)
        # Flag if: many rates but few/no matched rules
        # Threshold: more than 2 + 2*matched unaccounted hits
        if len(fee_hits) > matched * 2 + 2:
            pos, rate_str = fee_hits[0]
            start   = max(0, pos - 30)
            end     = min(len(page_text), pos + 80)
            snippet = re.sub(r"\s+", " ", page_text[start:end]).strip()
            gaps.append({
                "page":          page_num,
                "pct_found":     len(fee_hits),
                "rules_matched": matched,
                "snippet":       snippet,
            })

    return gaps[:8]


def _smart_truncate(text: str) -> str:
    """
    Keep the first HEAD_CHARS (agreement header: PSP name, entity, dates) +
    the last TAIL_CHARS (fee schedules almost always appear at the end).
    This approach is more reliable than a simple prefix truncation.
    """
    total = HEAD_CHARS + TAIL_CHARS
    if len(text) <= total:
        return text
    head = text[:HEAD_CHARS]
    tail = text[-TAIL_CHARS:]
    skipped = len(text) - total
    return (
        head
        + f"\n\n[... {skipped:,} characters of legal boilerplate omitted ...]\n\n"
        + tail
    )


# ── Post-processing normalisers ────────────────────────────────────────────

def _normalise_pct(value) -> float | None:
    """Ensure pct_rate values > 1 are divided by 100 (model sometimes returns 3.5 instead of 0.035)."""
    if isinstance(value, (int, float)) and value > 1:
        return round(value / 100, 8)
    return value


def _normalise_rule(rule: dict) -> dict:
    rule.setdefault("tiers", [])
    rule["pct_rate"] = _normalise_pct(rule.get("pct_rate"))
    for tier in rule.get("tiers", []):
        tier["pct_rate"] = _normalise_pct(tier.get("pct_rate")) or 0
    # Clamp pct_rate: > 0.99 is almost certainly wrong (model forgot to divide)
    if isinstance(rule.get("pct_rate"), float) and rule["pct_rate"] > 0.99:
        rule["pct_rate"] = round(rule["pct_rate"] / 100, 8)
    return rule


def _dedup_rules(rules: list) -> list:
    """
    Remove duplicate fee rules produced by the AI extracting the same fee twice
    under slightly different descriptions (e.g. 'Rwanda' vs 'New Jurisdiction Rwanda').

    Dedup key: (payment_method, fee_type, country, fee_kind, pct_rate, fixed_amount,
                fixed_currency) — i.e. identical fee economics regardless of description.
    When duplicates are found the first occurrence is kept (usually more specific description).
    """
    seen = {}
    out  = []
    for rule in rules:
        key = (
            (rule.get("payment_method") or "").lower().strip(),
            (rule.get("fee_type")       or "").lower().strip(),
            (rule.get("country")        or "").lower().strip(),
            (rule.get("fee_kind")       or ""),
            rule.get("pct_rate"),
            rule.get("fixed_amount"),
            (rule.get("fixed_currency") or ""),
        )
        if key not in seen:
            seen[key] = True
            out.append(rule)
    return out


def _validate_rules(rules: list) -> list[str]:
    """
    Validate extracted fee rules against allowed enums.
    Returns a list of human-readable warning strings (empty = all good).
    """
    warnings = []
    for i, rule in enumerate(rules):
        label = f"Rule {i+1}"
        ft = rule.get("fee_type")
        if ft not in FEE_TYPES:
            warnings.append(f"{label}: unknown fee_type '{ft}'")
        pm = rule.get("payment_method")
        if pm is not None and pm not in PAYMENT_METHODS:
            warnings.append(f"{label}: unknown payment_method '{pm}'")
        fk = rule.get("fee_kind")
        if fk not in ("percentage", "fixed", "fixed_plus_pct", "tiered"):
            warnings.append(f"{label}: unknown fee_kind '{fk}'")
        if fk in ("percentage", "fixed_plus_pct") and rule.get("pct_rate") is None:
            warnings.append(f"{label}: fee_kind='{fk}' but pct_rate is null")
        if fk in ("fixed", "fixed_plus_pct") and rule.get("fixed_amount") is None:
            warnings.append(f"{label}: fee_kind='{fk}' but fixed_amount is null")
    return warnings


# ── Main API call ──────────────────────────────────────────────────────────

def _call_claude(system_prompt: str, user_msg: str) -> str:
    """Call Claude and return the raw JSON text response."""
    if not ANTHROPIC_KEY:
        raise RuntimeError(
            "ANTHROPIC_API_KEY is not set. "
            "Add it to recon-app/.env: ANTHROPIC_API_KEY=sk-ant-..."
        )
    client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)
    message = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=4096,
        temperature=0.0,
        system=system_prompt,
        messages=[{"role": "user", "content": user_msg}],
    )
    raw = message.content[0].text.strip()
    # Strip markdown code fences if Claude wraps the JSON
    if raw.startswith("```"):
        raw = raw.split("```", 2)[1]          # drop opening fence + lang tag
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.rsplit("```", 1)[0].strip()
    return raw


def analyze_agreement(text: str, system_prompt: str = None,
                      pages: list = None) -> dict:
    """
    Send extracted agreement text to Claude and return a normalised dict:
      { "agreement": {...}, "fee_rules": [...] }
    """
    active_prompt = system_prompt or SYSTEM_PROMPT
    trimmed = _smart_truncate(text)

    user_msg = (
        "Parse the following PSP agreement document.\n\n"
        "Return a JSON object that EXACTLY matches this structure:\n"
        f"{_JSON_EXAMPLE}\n\n"
        "ALLOWED fee_type values (use ONLY these):\n"
        f"{FEE_TYPES}\n\n"
        "ALLOWED payment_method values (use ONLY these, or null):\n"
        f"{PAYMENT_METHODS}\n\n"
        "DOCUMENT TEXT FOLLOWS:\n"
        "─" * 60 + "\n"
        f"{trimmed}\n"
        "─" * 60 + "\n\n"
        "Return ONLY the JSON object. No markdown, no commentary."
    )

    raw    = _call_claude(active_prompt, user_msg)
    result = json.loads(raw)

    result.setdefault("agreement", {})
    result.setdefault("fee_rules", [])

    rules = [_normalise_rule(r) for r in result["fee_rules"]]
    before          = len(rules)
    rules           = _dedup_rules(rules)
    dups_removed    = before - len(rules)
    warnings        = _validate_rules(rules)

    if pages:
        annotate_sources(rules, pages)
    result["fee_rules"]     = rules
    result["dups_removed"]  = dups_removed
    result["warnings"]      = warnings
    result["gaps"]          = find_potential_gaps(rules, pages) if pages else []
    result["raw_response"]  = raw
    return result


# ── Amendment parsing ──────────────────────────────────────────────────────

AMENDMENT_SYSTEM_PROMPT = """\
You are a specialist financial-document parser. You are parsing an AMENDMENT to an \
existing PSP (Payment Service Provider) merchant agreement, NOT a full agreement. \
Your job is to extract ONLY the rules that are explicitly NEW, CHANGED, or REMOVED \
by this amendment, and return them as a single valid JSON object. \
No markdown fences, no explanation, no extra keys.

═══════════════════════════════════════════════════════════
  CRITICAL RULE — WHAT TO INCLUDE vs. EXCLUDE
═══════════════════════════════════════════════════════════
INCLUDE a rule ONLY if the amendment document explicitly:
  • States it is a NEW fee/market/payment method being added
  • States it REPLACES, SUPERSEDES, or CHANGES an existing fee
  • States an existing fee is REMOVED, CANCELLED, or DELETED

EXCLUDE (do NOT include) any rule that the amendment says:
  • "shall remain in place"
  • "remain unchanged"
  • "except as expressly provided … all other terms remain in full effect"
  • Is mentioned only in boilerplate, definitions, or legal clauses

═══════════════════════════════════════════════════════════
  OUTPUT JSON SCHEMA
═══════════════════════════════════════════════════════════
{
  "amendment": {
    "addendum_date": "<YYYY-MM-DD — the execution/effective date of this amendment, or null>",
    "notes":         "<one-sentence plain-English summary of what changed>"
  },
  "rule_changes": [
    {
      "action":       "<'add' | 'replace' | 'remove'>",
      "match_on": {
        "payment_method": "<value from PAYMENT METHODS list or null>",
        "country":        "<country name exactly as in amendment, or 'GLOBAL'>",
        "fee_type":       "<value from FEE TYPES list>"
      },
      "payment_method":  "<value from PAYMENT METHODS list, or null>",
      "fee_type":        "<value from FEE TYPES list>",
      "country":         "<country name exactly as written, or 'GLOBAL'>",
      "sub_provider":    "<card network / mobile operator, or null>",
      "fee_kind":        "<'percentage' | 'fixed' | 'fixed_plus_pct' | 'tiered'>",
      "pct_rate":        <decimal fraction or null — e.g. 2.5% → 0.025>,
      "fixed_amount":    <number or null>,
      "fixed_currency":  "<ISO-4217 code or null>",
      "description":     "<verbatim fee label or null>",
      "tiers":           []
    }
  ]
}

═══════════════════════════════════════════════════════════
  ACTION FIELD RULES
═══════════════════════════════════════════════════════════
"add"     — This is a NEW rule not previously covered (new country, new payment method,
            new fee type). match_on may be the same as the new rule values.
"replace" — An EXISTING rule with the same country + payment_method + fee_type is being
            given NEW rates. Set match_on to identify the old rule to overwrite.
"remove"  — An EXISTING rule is being explicitly deleted or cancelled.
            Set match_on to identify it. The fee fields (pct_rate, etc.) may be null.

═══════════════════════════════════════════════════════════
  PAYMENT METHODS (map document terms to these exactly)
═══════════════════════════════════════════════════════════
  "Credit Cards"        ← CREDIT CARDS, CARD, VISA, MASTERCARD, AMEX
  "Bank Wire"           ← BANK WIRE, WIRE TRANSFER, BANK TRANSFER, EFT, SPEI, PSE, PIX
  "Mobile Money"        ← MOBILE MONEY, MOMO, MOBILE PAYMENTS, MPESA, OPAY
  "Electronic Payment"  ← ELECTRONIC PAYMENT, E-WALLET (general), NETELLER, SKRILL
  "Crypto"              ← CRYPTO, CRYPTOCURRENCY, BITCOIN, USDT
  "MOMO"                ← use only when document explicitly labels it MOMO
  "E-Wallet"            ← use only when document labels it E-WALLET distinctly
  null                  ← when method is unspecified or applies to all

═══════════════════════════════════════════════════════════
  FEE TYPES (map document terms to these exactly)
═══════════════════════════════════════════════════════════
  "Deposit"          ← DEPOSIT FEE, ACQUIRING FEE, PROCESSING FEE, PAYINS
  "Withdrawal"       ← WITHDRAWAL FEE, CASHOUT FEE, PAYOUT FEE, PAYOUTS
  "Settlement"       ← SETTLEMENT FEE
  "Chargeback"       ← CHARGEBACK FEE, DISPUTE FEE
  "Refund"           ← REFUND FEE, REVERSAL FEE
  "Rolling Reserve"  ← ROLLING RESERVE, RESERVE RATE
  "Holdback"         ← HOLDBACK
  "Setup"            ← SETUP FEE, ONBOARDING FEE
  "Registration"     ← REGISTRATION FEE, ANNUAL CARD SCHEME FEE
  "Minimum Monthly"  ← MINIMUM FEE, MINIMUM MONTHLY FEE

  "PAYINS" in the document → fee_type = "Deposit"
  "PAYOUTS" in the document → fee_type = "Withdrawal"
  If a fee applies to both deposit AND withdrawal → create TWO rules.

═══════════════════════════════════════════════════════════
  RATE PARSING RULES
═══════════════════════════════════════════════════════════
  "2.5%"              → fee_kind="percentage", pct_rate=0.025
  "1.40% + 0.60 USD"  → fee_kind="fixed_plus_pct", pct_rate=0.014, fixed_amount=0.60, fixed_currency="USD"
  "3.10% + 0.50 USD"  → fee_kind="fixed_plus_pct", pct_rate=0.031, fixed_amount=0.50, fixed_currency="USD"
  Volume tiers (e.g. "0-300,000 USD: 5.70% / 300,001-600,000 USD: 5.60%"):
    → fee_kind="tiered", populate tiers array, pct_rate=null
  For tiered rules: volume_to=null for the last (open-ended) tier.
  NEVER invent a rate. If ambiguous, note it in description.

  Minimum fee language like "minimum fee of 0.20 USD" — capture in description field,
  do NOT create a separate fee rule for it.

═══════════════════════════════════════════════════════════
  AMENDMENT DATE
═══════════════════════════════════════════════════════════
  Look for: "Executed on", "Effective Date", "dated", signature block date.
  Convert any format to YYYY-MM-DD. If multiple dates, use the execution/signing date.

═══════════════════════════════════════════════════════════
  WORKED EXAMPLE (Blaven amendment, Feb 2025)
═══════════════════════════════════════════════════════════
The amendment adds PAYINS + PAYOUTS for Nigeria, Kenya, Ghana, South Africa, Brazil,
Mexico (tiered), Colombia (tiered), Ecuador (tiered), Peru.

Example output for two of those rules:
{
  "action": "add",
  "match_on": { "payment_method": "Mobile Money", "country": "Kenya", "fee_type": "Deposit" },
  "payment_method": "Mobile Money",
  "fee_type": "Deposit",
  "country": "Kenya",
  "sub_provider": null,
  "fee_kind": "percentage",
  "pct_rate": 0.03,
  "fixed_amount": null,
  "fixed_currency": null,
  "description": "PAYINS: Mobile Money - 3.00% (Flat Fee)",
  "tiers": []
}
{
  "action": "add",
  "match_on": { "payment_method": "Mobile Money", "country": "Kenya", "fee_type": "Withdrawal" },
  "payment_method": "Mobile Money",
  "fee_type": "Withdrawal",
  "country": "Kenya",
  "sub_provider": null,
  "fee_kind": "fixed_plus_pct",
  "pct_rate": 0.031,
  "fixed_amount": 0.50,
  "fixed_currency": "USD",
  "description": "PAYOUTS: 3.10% + 0.50 USD (Flat Fee)",
  "tiers": []
}
"""

# Amendment output schema (used when model supports structured outputs)
_AMENDMENT_SCHEMA = {
    "name":   "psp_amendment_extraction",
    "strict": True,
    "schema": {
        "type": "object",
        "required": ["amendment", "rule_changes"],
        "additionalProperties": False,
        "properties": {
            "amendment": {
                "type": "object",
                "required": ["addendum_date", "notes"],
                "additionalProperties": False,
                "properties": {
                    "addendum_date": {"type": ["string", "null"]},
                    "notes":         {"type": ["string", "null"]},
                },
            },
            "rule_changes": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": [
                        "action", "match_on",
                        "payment_method", "fee_type", "country", "sub_provider",
                        "fee_kind", "pct_rate", "fixed_amount", "fixed_currency",
                        "description", "tiers",
                    ],
                    "additionalProperties": False,
                    "properties": {
                        "action": {"type": "string", "enum": ["add", "replace", "remove"]},
                        "match_on": {
                            "type": "object",
                            "required": ["payment_method", "country", "fee_type"],
                            "additionalProperties": False,
                            "properties": {
                                "payment_method": {"type": ["string", "null"]},
                                "country":        {"type": "string"},
                                "fee_type":       {"type": "string"},
                            },
                        },
                        "payment_method":  {"type": ["string", "null"], "enum": PAYMENT_METHODS + [None]},
                        "fee_type":        {"type": "string", "enum": FEE_TYPES},
                        "country":         {"type": "string"},
                        "sub_provider":    {"type": ["string", "null"]},
                        "fee_kind":        {"type": "string", "enum": ["percentage", "fixed", "fixed_plus_pct", "tiered"]},
                        "pct_rate":        {"type": ["number", "null"]},
                        "fixed_amount":    {"type": ["number", "null"]},
                        "fixed_currency":  {"type": ["string", "null"]},
                        "description":     {"type": ["string", "null"]},
                        "tiers": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "required": ["volume_from", "volume_to", "pct_rate"],
                                "additionalProperties": False,
                                "properties": {
                                    "volume_from": {"type": "number"},
                                    "volume_to":   {"type": ["number", "null"]},
                                    "pct_rate":    {"type": "number"},
                                },
                            },
                        },
                    },
                },
            },
        },
    },
}


def analyze_amendment(text: str, system_prompt: str = None,
                      pages: list = None) -> dict:
    """
    Parse an amendment document and return:
      { "amendment": { "addendum_date": ..., "notes": ... },
        "rule_changes": [ { "action", "match_on", ...fee fields... }, ... ] }
    """
    active_prompt = system_prompt or AMENDMENT_SYSTEM_PROMPT
    trimmed = _smart_truncate(text)

    user_msg = (
        "Parse the following PSP AMENDMENT document.\n\n"
        "Return a JSON object with keys 'amendment' and 'rule_changes'.\n\n"
        "ALLOWED fee_type values (use ONLY these):\n"
        f"{FEE_TYPES}\n\n"
        "ALLOWED payment_method values (use ONLY these, or null):\n"
        f"{PAYMENT_METHODS}\n\n"
        "DOCUMENT TEXT FOLLOWS:\n"
        + "─" * 60 + "\n"
        + trimmed + "\n"
        + "─" * 60 + "\n\n"
        "Return ONLY the JSON object. No markdown, no commentary."
    )

    raw    = _call_claude(active_prompt, user_msg)
    result = json.loads(raw)

    result.setdefault("amendment", {})
    result.setdefault("rule_changes", [])

    rules        = [_normalise_rule(r) for r in result["rule_changes"]]
    before       = len(rules)
    rules        = _dedup_rules(rules)
    warnings     = _validate_rules(rules)

    if pages:
        annotate_sources(rules, pages)
    result["rule_changes"]  = rules
    result["dups_removed"]  = before - len(rules)
    result["warnings"]      = warnings
    result["gaps"]          = find_potential_gaps(rules, pages) if pages else []
    result["raw_response"]  = raw
    return result
